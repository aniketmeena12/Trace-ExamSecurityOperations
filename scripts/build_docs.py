"""Generate a comprehensive Word (.docx) document describing the Trace project.

Run:  backend/.venv/bin/python scripts/build_docs.py
Output: Trace_Project_Documentation.docx in the repo root.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette (matches the app's "security ops center" theme) -----------------
INK      = RGBColor(0x14, 0x1B, 0x2E)   # near-black navy
ACCENT   = RGBColor(0x0E, 0x7A, 0x8C)   # deep cyan
ACCENT2  = RGBColor(0x1B, 0xA5, 0xB8)   # bright cyan
DANGER   = RGBColor(0xC0, 0x2A, 0x45)
VERIFY   = RGBColor(0x1E, 0x8E, 0x5A)
MUTE     = RGBColor(0x5B, 0x66, 0x77)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, hex_fill):
    """Set a table cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1BA5B8")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def code_block(doc, text):
    """Monospace, shaded paragraph for code/commands."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    # shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F5F7")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = INK
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = ACCENT
        p.add_run("  " + text)
    else:
        p.add_run(text)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level == 1 else INK
    return h


def para(doc, text, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def make_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=WHITE, size=10)
        shade(hdr[i], "0E7A8C")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


# =============================================================================
doc = Document()

# global style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK

# ---- TITLE PAGE -------------------------------------------------------------
for _ in range(3):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("TRACE")
r.font.size = Pt(54)
r.bold = True
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Leak-Proof, Traceable Exam Papers")
r.font.size = Pt(20)
r.font.color.rgb = INK

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("Make an exam paper impossible to open before the exam —\nand traceable to its exact source if it ever leaks.")
r.font.size = Pt(12)
r.italic = True
r.font.color.rgb = MUTE

doc.add_paragraph()
add_hr(doc)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("India's FAR AWAY 2026  ·  Theme: Secure, Leak-Proof Examinations")
r.font.size = Pt(11)
r.bold = True
r.font.color.rgb = ACCENT2

stack = doc.add_paragraph()
stack.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = stack.add_run("Python · FastAPI · SQLite · React · Vite · Tailwind\nAES-256-GCM · Shamir 3-of-5 (from scratch) · DCT watermarking · SHA-256 hash chain")
r.font.size = Pt(10)
r.font.color.rgb = MUTE

doc.add_page_break()

# ---- TABLE OF CONTENTS ------------------------------------------------------
heading(doc, "Contents", 1)
toc_items = [
    "1.  Executive Summary — what Trace is, in one minute",
    "2.  The Problem It Solves",
    "3.  The Three Guarantees",
    "4.  How It Works, End to End (the demo story)",
    "5.  System Architecture",
    "6.  Technology Stack",
    "7.  Guarantee 1 — Multi-Custodian, Time-Locked Decryption",
    "8.  Guarantee 2 — Invisible Per-Candidate Watermarking",
    "9.  Guarantee 3 — Tamper-Evident Audit Log",
    "10. Defense in Depth — Dynamic Assembly & Leak-Match Detector",
    "11. The Backend — Modules & Data Model",
    "12. The API — Every Endpoint",
    "13. The Frontend — Four Role Dashboards",
    "14. Security Model & Trust Boundary",
    "15. Testing & Verification",
    "16. How to Run It Yourself",
    "17. Build Status (Milestones)",
    "18. Project Layout",
    "19. Glossary",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ---- 1. EXECUTIVE SUMMARY ---------------------------------------------------
heading(doc, "1.  Executive Summary", 1)
para(doc,
    "Exam paper leaks are a recurring, high-impact failure of trust. A single insider with "
    "early access can compromise an entire exam for millions of candidates. Trace is a working "
    "system that removes that possibility and, if a leak somehow happens through a screen photo "
    "or printout, traces the leaked copy back to the exact candidate it was issued to.")
para(doc,
    "Trace rests on three real security guarantees — implemented in code and proven by tests, "
    "not slideware:")
bullet(doc, "the paper cannot be opened by any single person, and cannot be opened before the exam start time;",
       bold_lead="Sealed until the moment:")
bullet(doc, "every printed/displayed copy carries an invisible fingerprint identifying its candidate;",
       bold_lead="Traceable forever:")
bullet(doc, "every access is recorded in a log that cannot be silently altered.",
       bold_lead="Tamper-evident:")
para(doc,
    "The system is a full-stack application: a Python/FastAPI backend implementing the "
    "cryptography from first principles, and a React dashboard that lets the four real-world "
    "roles — Exam Controller, Custodian, Candidate, and Investigator — each play their part in a "
    "live, end-to-end demonstration.")

# ---- 2. THE PROBLEM ---------------------------------------------------------
heading(doc, "2.  The Problem It Solves", 1)
para(doc,
    "High-stakes exams (entrance tests, recruitment, board exams) depend on the paper staying "
    "secret until the exam begins. In practice two things go wrong:")
bullet(doc, "Somewhere between the printer and the exam hall, one trusted person opens the "
            "paper early and sells or shares it.", bold_lead="The insider leak —")
bullet(doc, "Even with secrecy, once a copy is out there is no way to know whose copy it was, "
            "so nobody is ever held accountable.", bold_lead="The accountability gap —")
para(doc,
    "Trace attacks both: it makes early opening cryptographically impossible (no single insider "
    "holds the key, and a server clock blocks decryption until release time), and it makes every "
    "copy individually identifiable so a leak names its own source.")

# ---- 3. THE THREE GUARANTEES ------------------------------------------------
heading(doc, "3.  The Three Guarantees", 1)
make_table(doc,
    ["Guarantee", "What it means", "How it is achieved"],
    [
        ["1. Just-in-time, multi-custodian decryption",
         "No single insider — and no pair — can open the paper, and nobody can open it before exam start.",
         "AES-256-GCM encryption; the key is split with Shamir's Secret Sharing (3-of-5) across independent custodians, plus a server-enforced release-time gate."],
        ["2. Per-candidate invisible watermarking",
         "Any leaked image traces back to its exact source candidate.",
         "An invisible DCT-domain fingerprint encoding center ID + candidate ID is embedded in each rendered copy; it survives JPEG recompression."],
        ["3. Tamper-evident audit log",
         "Altering any past event is detectable.",
         "An append-only SHA-256 hash-chained log of every access event; changing any entry breaks the chain."],
    ],
    widths=[1.7, 2.3, 2.6])

# ---- 4. HOW IT WORKS END TO END ---------------------------------------------
heading(doc, "4.  How It Works, End to End", 1)
para(doc, "The whole system can be experienced as one continuous story, with each role acting in turn:")

heading(doc, "Act 1 — Seal (the Exam Controller / Admin)", 2)
para(doc, "The controller uploads a paper and seals an exam. Behind the scenes the backend "
          "generates a fresh 256-bit key, encrypts the paper with AES-256-GCM, then splits the key "
          "into 5 custodian shares using Shamir's Secret Sharing — any 3 of which can rebuild it. "
          "The plaintext key is never stored. A release time is set (e.g. \"+30 seconds\" for the demo).")

heading(doc, "Act 2 — Authorize (the five Custodians)", 2)
para(doc, "Each custodian sees their own key share — masked until the release window opens. When the "
          "window opens, each submits their share. The dashboard shows a live quorum: \"2 of 3 "
          "required shares.\" Crucially, even when all three valid shares are present, the server "
          "refuses to decrypt until the clock passes release time.")

heading(doc, "Act 3 — Unlock (the time gate releases)", 2)
para(doc, "Once both conditions hold — enough shares AND release time reached — the server "
          "reconstructs the key by Lagrange interpolation, decrypts the paper (the GCM tag proves "
          "the reconstruction was correct), and re-wraps the key under a server secret so candidates "
          "can be served without the raw key ever leaving the server. The vault opens, live.")

heading(doc, "Act 4 — Read (the Candidate)", 2)
para(doc, "Before release the candidate sees only a countdown. After unlock they receive their own "
          "copy of the paper — rendered to an image and stamped with an invisible watermark unique to "
          "them. To the eye it is identical to everyone else's (PSNR ~36 dB); under the hood it carries "
          "their fingerprint.")

heading(doc, "Act 5 — Trace & Verify (the Investigator)", 2)
para(doc, "If a copy leaks, the investigator uploads the leaked image — even a re-compressed JPEG "
          "screen photo. Trace extracts the fingerprint and matches it against the issue registry, "
          "naming the exact candidate with a confidence score. Finally, the investigator clicks "
          "\"Verify Integrity\" to recompute the entire SHA-256 audit chain and confirm no record "
          "was tampered with.")

# ---- 5. ARCHITECTURE --------------------------------------------------------
heading(doc, "5.  System Architecture", 1)
para(doc, "The data flow, from sealed paper to forensic trace:")
code_block(doc,
    "        custodians (3-of-5 shares)         server release-time gate\n"
    "                   \\                       /\n"
    "  exam paper --AES-256-GCM--> encrypted vault --> JIT decrypt --> render\n"
    "                                    |                                |\n"
    "                                    v                      per-candidate DCT\n"
    "                          hash-chained audit log               watermark\n"
    "                                                                    |\n"
    "                                              leaked image --> forensic trace")
para(doc, "The backend is the single trusted enforcement point: it holds custodian shares in escrow, "
          "enforces the time gate, performs reconstruction and decryption, embeds watermarks, and "
          "maintains the audit chain. The frontend is a thin client — every animation it shows is "
          "driven by a real backend event, never mock data.")

# ---- 6. TECH STACK ----------------------------------------------------------
heading(doc, "6.  Technology Stack", 1)
make_table(doc,
    ["Layer", "Technology", "Role"],
    [
        ["Backend framework", "Python 3.10+ · FastAPI", "REST API, request validation, OpenAPI docs"],
        ["Database / ORM", "SQLite · SQLAlchemy 2.0", "Exams, users, shares, watermark registry, audit events"],
        ["Cryptography", "pycryptodome (AES-GCM) + from-scratch GF(256) Shamir", "Encryption and secret sharing"],
        ["Auth", "PyJWT (HS256) · OAuth2 password flow", "4 roles: admin / custodian / candidate / investigator"],
        ["Watermarking", "OpenCV (DCT) · NumPy · Pillow", "Invisible fingerprint embed / extract / trace"],
        ["Frontend", "React 18 · Vite 5 · Tailwind 3", "Four role dashboards, dark mission-control theme"],
        ["Frontend data/UX", "TanStack React Query 5 · Framer Motion 11", "Live polling of real endpoints; animations"],
        ["Testing", "pytest", "~816 lines across 8 test files, full crypto + API coverage"],
    ],
    widths=[1.5, 2.3, 2.8])

# ---- 7. GUARANTEE 1 ---------------------------------------------------------
heading(doc, "7.  Guarantee 1 — Multi-Custodian, Time-Locked Decryption", 1)
heading(doc, "AES-256-GCM encryption", 2)
bullet(doc, "256-bit (32-byte) key, generated fresh per exam with os.urandom.", bold_lead="Key:")
bullet(doc, "96-bit (12-byte) nonce — the GCM-recommended size.", bold_lead="Nonce:")
bullet(doc, "128-bit (16-byte) authentication tag — detects any tampering with the ciphertext.", bold_lead="Tag:")
bullet(doc, "Associated data binds the ciphertext to its context: \"center={id};exam={name}\".", bold_lead="AAD:")
para(doc, "Stored layout: nonce (12) || tag (16) || ciphertext. The plaintext key is never written to disk.")

heading(doc, "Shamir's Secret Sharing over GF(256) — built from scratch", 2)
para(doc, "Rather than calling a black-box library, the secret sharing is implemented from first "
          "principles over the finite field GF(2^8), the same field AES itself uses:")
bullet(doc, "The AES irreducible polynomial 0x11B (x^8 + x^4 + x^3 + x + 1) with generator 0x03.", bold_lead="Field:")
bullet(doc, "Addition is XOR; multiplication, division and inversion are O(1) via precomputed "
            "512-entry exp and 256-entry log tables.", bold_lead="Arithmetic:")
bullet(doc, "The 32-byte key is split byte-by-byte. Each byte becomes the constant term of a "
            "random degree k-1 polynomial; shares are evaluations f(x) at x = 1..n.", bold_lead="Split:")
bullet(doc, "Any k shares rebuild each secret byte via Lagrange interpolation at x = 0. In GF(256) "
            "negation is identity, so subtraction is also XOR.", bold_lead="Reconstruct:")
bullet(doc, "k shares always reconstruct; k-1 shares leave the secret uniformly random — "
            "information-theoretically secure, not just computationally hard.", bold_lead="Security:")
para(doc, "The demo default is 3-of-5: no single custodian and no pair can rebuild the key; it takes "
          "three distinct custodians authorizing together.")

heading(doc, "The server-enforced time gate", 2)
para(doc, "Reconstruction is gated by two independent guards. Even when all three valid shares are "
          "submitted, the server refuses to decrypt until the wall-clock passes release_time. This "
          "guard is unconditional — there is no override — and every blocked attempt is logged as an "
          "UNLOCK_DENIED audit event with the seconds remaining.")

# ---- 8. GUARANTEE 2 ---------------------------------------------------------
heading(doc, "8.  Guarantee 2 — Invisible Per-Candidate Watermarking", 1)
para(doc, "Each candidate's paper is rendered to an image and stamped with an invisible fingerprint "
          "in the frequency domain, using a Koch–Zhao differential DCT scheme.")
heading(doc, "What is embedded", 2)
bullet(doc, "A 16-bit MAGIC marker (0xACE5) used to detect whether a watermark is present at all.")
bullet(doc, "A 64-bit fingerprint = BLAKE2b-64 of (center ID, exam:candidate code) — unique per "
            "candidate per exam. Total payload: 80 bits.")
heading(doc, "How it is embedded", 2)
bullet(doc, "The image luminance is split into 8x8 DCT blocks (the same block size JPEG uses).")
bullet(doc, "Two mid-band coefficients at positions (2,3) and (3,2) are nudged so their relative "
            "order encodes one payload bit, with an enforced gap of 26 between them.")
bullet(doc, "Payload bits are repeated across hundreds of blocks, then an inverse DCT rebuilds the image.")
para(doc, "The result is imperceptible to the eye — PSNR around 36 dB — yet robust.")
heading(doc, "Why it survives a leak", 2)
para(doc, "A real leak is rarely the original file — it is a screen photo or a re-saved JPEG. The "
          "watermark survives because it encodes the relative order of mid-band coefficients (which "
          "JPEG quantization preserves far better than absolute values) and because each bit is "
          "spread across many blocks, so a few destroyed blocks cannot flip the aggregate decision. "
          "Extraction uses a soft decision: it sums the coefficient differences across all blocks and "
          "reads the sign.")
heading(doc, "Tracing a leak", 2)
para(doc, "The investigator's upload is run through extraction, and the recovered 64-bit fingerprint "
          "is matched against the registry of issued watermarks by Hamming distance (tolerating up to "
          "2 flipped magic bits). The output reports whether a watermark was present, the matched "
          "candidate, the bit distance, and a confidence score = 1 - (bit_distance / 64).")

# ---- 9. GUARANTEE 3 ---------------------------------------------------------
heading(doc, "9.  Guarantee 3 — Tamper-Evident Audit Log", 1)
para(doc, "Every meaningful action is recorded as an event in an append-only log, and each event is "
          "chained to the one before it by a SHA-256 hash — the same idea that makes a blockchain "
          "tamper-evident.")
para(doc, "Each event's hash is computed over its own fields plus the previous event's hash:")
code_block(doc, "hash(i) = SHA256( id | timestamp | actor | action | target | details | prev_hash )")
bullet(doc, "The chain starts from a fixed genesis hash (64 zeros).")
bullet(doc, "Recorded actions include: LOGIN, EXAM_SEALED, SHARE_SUBMITTED, UNLOCK_DENIED, "
            "PAPER_UNLOCKED, PAPER_ACCESSED, WATERMARK_ISSUED, LEAK_TRACED, AUDIT_VERIFIED, "
            "QUESTION_ADDED, PAPER_ASSEMBLED, and LEAK_MATCHED.")
bullet(doc, "Verification recomputes every hash and checks every prev_hash link. If anyone edits, "
            "inserts, or deletes a row, the recomputed hash stops matching and the break is reported "
            "with the exact event id.")
bullet(doc, "A threading lock serialises appends so concurrent writes cannot corrupt the tail "
            "(a real concurrency race that was caught and fixed during development).")

# ---- 10. DEFENSE IN DEPTH ---------------------------------------------------
heading(doc, "10.  Defense in Depth — Dynamic Assembly & Leak-Match Detector", 1)
para(doc, "Beyond the three core guarantees, Trace adds two further layers that directly shrink "
          "the chance and the blast radius of a leak — and trace one even when no image exists.")

heading(doc, "10.1  Dynamic per-candidate paper assembly", 2)
para(doc, "A static exam is one secret blob: leak it once and the whole exam is compromised, for "
          "everyone. Dynamic assembly removes that single point of failure. Instead of one paper, the "
          "system keeps a large bank of individually-encrypted questions; at release time a blueprint "
          "(sections + counts) and a per-candidate seed select a unique subset and assemble that "
          "candidate's paper on demand.")
bullet(doc, "Each question body is AES-256-GCM ciphertext at rest, under a server question-bank key "
            "that is domain-separated from the vault key — a database dump cannot read the bank.",
       bold_lead="Encrypted bank —")
bullet(doc, "The selection seed is HMAC(exam_key, \"exam:candidate:section\"). Because the exam key "
            "only exists after the Shamir quorum AND the time gate open, no setter, admin, or "
            "database reader can predict which questions reach which candidate before release.",
       bold_lead="Unpredictable —")
bullet(doc, "Every candidate's paper is content-distinct, so a leaked copy exposes only one variant — "
            "and, combined with the DCT watermark, still names its owner.",
       bold_lead="Low blast radius —")
bullet(doc, "At seal time the per-section pools are frozen inside the encrypted manifest, so retiring "
            "or editing a bank question later cannot change (or break) an already-sealed exam's papers.",
       bold_lead="Immutable once sealed —")
bullet(doc, "Only the selection (a list of question ids) is recorded per candidate — the assembled "
            "paper is never persisted in plaintext; it exists only in memory while one request is served.",
       bold_lead="Never stored —")

heading(doc, "10.2  Leak-match detector (text-only leaks)", 2)
para(doc, "Real leaks are often just text — a question pasted into a chat group, with no image to "
          "trace. The leak-match detector handles exactly this case. An investigator pastes the "
          "suspected text and Trace:")
bullet(doc, "matches it against the encrypted bank by token-containment on each question's prompt "
            "(option labels are excluded so short prompts are not over-matched), flagging every "
            "question the text contains;", bold_lead="Identifies what leaked —")
bullet(doc, "intersects the per-candidate selection records: a candidate whose paper contained every "
            "matched question is a prime suspect; otherwise the highest-overlap candidate is the "
            "strongest lead. The issued-watermark fingerprint is attached for an image follow-up.",
       bold_lead="Narrows who leaked it —")
para(doc, "This is where dynamic assembly pays off forensically: because every candidate received a "
          "different combination, a leak containing several questions intersects to a very small "
          "suspect set — often a single candidate.")

# ---- 11. BACKEND ------------------------------------------------------------
heading(doc, "11.  The Backend — Modules & Data Model", 1)
heading(doc, "Module map", 2)
make_table(doc,
    ["Module", "Responsibility"],
    [
        ["crypto/", "gf256, shamir, aes_gcm — the from-scratch cryptographic core"],
        ["security/", "passwords (bcrypt), tokens (JWT), deps (role guards), keywrap (server key wrapping)"],
        ["audit/", "chain.py — SHA-256 hash-chained event log + verification"],
        ["watermark/", "core (DCT embed/extract), render (text -> image), trace (registry match)"],
        ["services/", "exams.py (seal/unlock/serve), assembly.py (dynamic question bank), leakmatch.py (text leak detector), watermarking.py"],
        ["api/", "app.py (FastAPI factory) + routers/ (auth, exams, questions, investigator, audit)"],
        ["models.py / schemas.py / db.py", "SQLAlchemy tables, Pydantic request/response shapes, DB session"],
        ["bootstrap.py / config.py", "demo seeding + idempotent migration; env-driven settings"],
    ],
    widths=[2.0, 4.5])

heading(doc, "Data model (tables)", 2)
make_table(doc,
    ["Table", "Key fields", "Purpose"],
    [
        ["users", "username, role, password_hash, candidate_code, center_id", "Accounts for the four roles"],
        ["exams", "name, release_time, threshold_k, nonce, tag, ciphertext, aad, status, released_key_wrapped, assembly_mode, blueprint", "A sealed paper/manifest and its unlock state"],
        ["custodian_shares", "exam_id, custodian_id, x, y", "Each custodian's escrowed Shamir share"],
        ["share_submissions", "exam_id, custodian_id, submitted_at", "Authorizations during the unlock ceremony"],
        ["questions", "subject, section, topic, difficulty, nonce, tag, ciphertext, contributor", "The encrypted question bank (bodies are ciphertext)"],
        ["candidate_papers", "exam_id, username, selected_question_ids, assembled_at", "Per-candidate selection (ids only, never plaintext)"],
        ["issued_watermarks", "exam_id, username, fingerprint, center_id, candidate_code", "Registry for tracing leaked images"],
        ["audit_events", "timestamp, actor, action, target, details, prev_hash, hash", "The tamper-evident hash chain"],
    ],
    widths=[1.4, 3.0, 2.1])

# ---- 12. API ----------------------------------------------------------------
heading(doc, "12.  The API — Every Endpoint", 1)
make_table(doc,
    ["Method", "Path", "Role", "Purpose"],
    [
        ["POST", "/auth/login", "any", "OAuth2 password login -> JWT"],
        ["GET", "/auth/me", "any", "Current user profile"],
        ["POST", "/exams", "admin", "Seal a paper/manifest (encrypt + Shamir-split the key)"],
        ["GET", "/exams", "any", "List all exams"],
        ["GET", "/exams/{id}", "any", "Inspect one exam's metadata"],
        ["GET", "/exams/{id}/unlock/status", "any", "Live unlock progress (drives the vault UI)"],
        ["GET", "/exams/{id}/blueprint", "any", "Dynamic exam structure (counts, not content)"],
        ["GET", "/exams/{id}/my-share", "custodian", "Fetch own share (masked until window opens)"],
        ["POST", "/exams/{id}/shares/submit", "custodian", "Authorize release of one share"],
        ["GET", "/exams/{id}/paper", "candidate/admin/investigator", "Read released paper (assembled per-candidate if dynamic)"],
        ["GET", "/exams/{id}/paper/image", "candidate/admin", "Per-candidate watermarked PNG"],
        ["POST", "/questions", "admin", "Add one AES-GCM-encrypted question to the bank"],
        ["GET", "/questions", "admin", "List bank metadata only (never bodies)"],
        ["POST", "/investigator/trace", "investigator/admin", "Trace a leaked image to a candidate"],
        ["POST", "/investigator/match", "investigator/admin", "Match leaked text to bank + narrow the source"],
        ["GET", "/audit", "investigator/admin", "The append-only event log"],
        ["GET", "/audit/verify", "investigator/admin", "Recompute the hash chain -> intact?"],
    ],
    widths=[0.7, 2.4, 1.5, 1.9])
para(doc, "Authentication is a JWT bearer token (HS256) issued at login, carrying the username and "
          "role, with a configurable lifetime (default 12 hours). Every protected route checks the "
          "role; missing/invalid tokens get 401, insufficient role gets 403.")

# ---- 13. FRONTEND -----------------------------------------------------------
heading(doc, "13.  The Frontend — Four Role Dashboards", 1)
para(doc, "A single dark \"security operations center\" web app hosts all four roles. A floating "
          "role-switcher (bottom center) jumps between them instantly — each identity is a real JWT "
          "session pre-authenticated against the backend, so there is no mock data anywhere.")
heading(doc, "Admin (Exam Controller)", 2)
para(doc, "Seal an exam (name, subject, center, threshold, release window), browse sealed exams, and "
          "watch a live operations panel: the vault animation, the share quorum, the encryption "
          "badge, and the time-gate countdown — all polled from the real backend every 1.5 seconds. "
          "A Static/Dynamic toggle switches the seal form to a blueprint editor for dynamic exams, and "
          "the operations panel shows a live \"dynamic assembly\" strip (bank pool + per-candidate count).")
heading(doc, "Custodian", 2)
para(doc, "See your own key share (rendered as a masked pixel grid until the window opens, then "
          "revealed as hex), submit it with one action, and watch the shared quorum fill toward the "
          "3-of-5 threshold.")
heading(doc, "Candidate", 2)
para(doc, "Before release: a countdown or an \"awaiting shares\" progress bar. After unlock: your own "
          "invisibly-watermarked paper image, with your fingerprint shown and a download button.")
heading(doc, "Investigator", 2)
para(doc, "Three tools in one console: a Leak-Match Detector (paste suspected leaked text to see the "
          "matched bank questions with containment scores and a ranked suspect list); a Forensic Trace "
          "(drag-drop a leaked image to trace it to a candidate, with confidence bar and bit distance); "
          "and the Audit Ledger — click \"Verify Integrity\" to recompute the SHA-256 chain and confirm "
          "it is intact, with any broken events highlighted in red.")
heading(doc, "Design system", 2)
para(doc, "A small set of reusable primitives (StatusPill, MonoReadout, Card, Metric, CountdownTimer, "
          "ShareSlots, VaultState, LogView) give the app a consistent mission-control look: cyan for "
          "secure, amber for pending/time-locked, red for denied/leak, green for verified, with glass "
          "morphism surfaces and JetBrains Mono / Space Grotesk typography.")

# ---- 14. SECURITY MODEL -----------------------------------------------------
heading(doc, "14.  Security Model & Trust Boundary", 1)
para(doc, "Trace defends against the two threats that cause real exam leaks:")
bullet(doc, "No one custodian (or pair) can open a paper. The key only exists after k distinct "
            "custodians authorize AND the server reconstructs it via Shamir. The key is never stored whole.",
       bold_lead="A single malicious insider —")
bullet(doc, "The release-time gate is enforced server-side and is unconditional: even with every "
            "valid share present, decryption is refused until release_time.",
       bold_lead="Early opening —")
para(doc, "In this single-node demo the server is the trusted enforcement point and holds custodian "
          "shares in escrow; the released key is wrapped under a server secret (derived from the app "
          "secret, not stored raw) so a raw database dump cannot decrypt a sealed paper. Hardening "
          "against a fully compromised server — HSM-held or client-custodied shares — is called out "
          "honestly as future work rather than faked.")

heading(doc, "Threat Model: Direct Database Access", 2)
para(doc, "A natural question is: what happens if an attacker gains direct read access to the "
          "database and tries to \"open the vault\" there? The answer separates the encrypted vault "
          "from the escrowed shares.")
bullet(doc, "Reading the exams table yields only the AES-256-GCM encrypted paper (nonce, tag, and "
            "ciphertext). The plaintext paper and the AES key are never stored in the vault.",
       bold_lead="The vault stays sealed —")
bullet(doc, "Before release, no exam decryption key exists in the database at all. After a legitimate "
            "unlock, the released key is stored only in wrapped form, encrypted under a server-held "
            "secret derived from TRACE_SECRET_KEY that is maintained outside the database (an "
            "environment variable). Consequently, a database dump alone is insufficient to recover "
            "exam plaintext in either state.",
       bold_lead="No usable key in either state —")
bullet(doc, "AES-GCM authentication ensures that any modification of stored ciphertext is detected "
            "during decryption, so a forged or altered vault is rejected rather than silently served.",
       bold_lead="Tampering is detected —")
bullet(doc, "The primary remaining risk lies in the escrowed Shamir shares stored in the "
            "custodian_shares table: an attacker with unrestricted access to the entire database could "
            "recover enough shares to reconstruct the key, bypassing the application-layer quorum and "
            "release-time controls. Mitigating this requires moving shares into external custodial "
            "systems, HSMs, or trusted execution environments rather than relying on application-level "
            "controls alone.",
       bold_lead="The remaining soft spot —")
para(doc, "In summary, direct access to the encrypted vault does not reveal exam content — this is a "
          "cryptographic guarantee, not an application-level one. The remaining dependency is on "
          "trusted server-side share storage, which the production hardening above removes: under "
          "external, HSM-, or enclave-held shares, a database compromise alone would not be enough to "
          "recover an exam paper.")

# ---- 15. TESTING ------------------------------------------------------------
heading(doc, "15.  Testing & Verification", 1)
para(doc, "Correctness is proven by an automated test suite — 81 tests passing — covering the parts "
          "that matter most:")
make_table(doc,
    ["Area", "What is proven"],
    [
        ["GF(256) field", "Field laws (add, multiply, divide, invert), generator order"],
        ["Shamir 3-of-5", "Every one of the 10 valid triples reconstructs; no pair ever does; a single share leaks nothing"],
        ["AES-256-GCM", "Round-trip, tamper rejection, AAD binding, key/nonce/tag lengths"],
        ["Watermark", "DCT embed/extract, PSNR, magic detection, Hamming matching"],
        ["Audit chain", "Linkage, verification, tamper detection"],
        ["Dynamic assembly", "Determinism, distinctness across candidates, blueprint counts, no plaintext at rest, release gating, post-seal immutability"],
        ["Leak-match", "Exact-text match, no-match on unrelated text, owner-is-strongest-lead, role enforcement, audit"],
        ["API flow", "Role enforcement, the time gate blocking 3 valid shares, legitimate unlock, candidate access, watermark issue & trace"],
    ],
    widths=[1.6, 4.9])
para(doc, "Run them with:  cd backend && pytest -v")

# ---- 16. HOW TO RUN ---------------------------------------------------------
heading(doc, "16.  How to Run It Yourself", 1)
heading(doc, "Backend (API server)", 2)
code_block(doc,
    "cd backend\n"
    "python3 -m venv .venv\n"
    "source .venv/bin/activate\n"
    "pip install -r requirements.txt\n"
    "uvicorn trace.api.app:app --reload      # http://127.0.0.1:8000\n"
    "# interactive API docs at http://127.0.0.1:8000/docs")
para(doc, "On first start the database is created and demo accounts are seeded (dev-only passwords): "
          "admin/admin123, investigator/invest123, cust1..cust5 / custodian1..5, cand001/candidate1.")
heading(doc, "Frontend (dashboards)", 2)
code_block(doc,
    "cd frontend\n"
    "npm install\n"
    "npm run dev        # http://localhost:5173  (proxies /api -> backend:8000)")
heading(doc, "Command-line demos (no browser needed)", 2)
code_block(doc,
    "cd backend\n"
    "python -m cli.crypto_demo     # encrypt -> split -> 3-of-5 reconstruct, 2 blocked\n"
    "python -m cli.api_demo        # time gate holds 3 shares; release; audit tamper caught\n"
    "python -m cli.watermark_demo  # watermark, JPEG-leak, trace back to the candidate")
para(doc, "The demo loop in the UI: as Admin seal an exam set to \"+30 sec\"; as each Custodian submit "
          "a share; when 3 are in and the timer hits zero the vault opens; as the Candidate view and "
          "download your watermarked paper; as the Investigator upload that file to trace it, then "
          "Verify Integrity on the ledger.")

# ---- 17. MILESTONES ---------------------------------------------------------
heading(doc, "17.  Build Status (Milestones)", 1)
make_table(doc,
    ["Milestone", "Scope", "Status"],
    [
        ["M1", "Repo scaffold + crypto core (AES-256-GCM, Shamir 3-of-5, tests, CLI)", "Done"],
        ["M2", "Release-time gate, hash-chained audit log, JWT roles, FastAPI endpoints", "Done"],
        ["M3", "DCT watermark embed/extract/trace + JPEG-robustness test", "Done"],
        ["M3.5", "Watermark API wiring (image + trace + registry + custodian share)", "Done"],
        ["M4", "Frontend scaffold + 4 role dashboards wired to backend", "Done"],
        ["Dynamic assembly", "Encrypted question bank + per-candidate paper assembly (leak-resistance layer)", "Done"],
        ["Leak-match", "Text-leak detector: match to bank + narrow the source", "Done"],
        ["M5", "Hero visuals: forensic leak-trace reveal + time-lock vault", "In progress"],
        ["M6", "Seed data, one-command demo script, README, architecture diagram", "In progress"],
    ],
    widths=[1.3, 4.0, 1.2])

# ---- 18. PROJECT LAYOUT -----------------------------------------------------
heading(doc, "18.  Project Layout", 1)
code_block(doc,
    "trace/\n"
    "|- backend/\n"
    "|  |- trace/\n"
    "|  |  |- crypto/     gf256 . shamir . aes_gcm        <- crypto core\n"
    "|  |  |- security/   passwords . tokens(JWT) . deps . keywrap\n"
    "|  |  |- audit/      chain.py  (SHA-256 hash chain)\n"
    "|  |  |- watermark/  core(DCT) . render . trace      <- watermarking\n"
    "|  |  |- services/   exams.py . assembly.py (dynamic bank) . leakmatch.py . watermarking.py\n"
    "|  |  |- api/        app.py + routers/ (auth . exams . questions . investigator . audit)\n"
    "|  |  |- models.py . schemas.py . db.py . config.py . bootstrap.py\n"
    "|  |- tests/         gf256 . shamir . aes_gcm . watermark . audit_chain . api . assembly . leakmatch\n"
    "|  |- cli/           crypto_demo.py . api_demo.py . watermark_demo.py\n"
    "|- frontend/         React app (Vite + Tailwind + React Query)\n"
    "|  |- src/\n"
    "|     |- components/ design system + LogView, CountdownTimer, ShareSlots, VaultState\n"
    "|     |- dashboards/ Admin . Custodian . Candidate . Investigator\n"
    "|     |- api/        client.js . hooks.js (React Query, real endpoints)\n"
    "|     |- auth/       AuthContext (auto-login + role switching)\n"
    "|- sample_data/      sample_paper.txt\n"
    "|- scripts/          run_demo.sh, build_docs.py")

# ---- 19. GLOSSARY -----------------------------------------------------------
heading(doc, "19.  Glossary", 1)
make_table(doc,
    ["Term", "Plain-language meaning"],
    [
        ["AES-256-GCM", "A strong, authenticated encryption algorithm. \"Authenticated\" means any tampering with the encrypted data is detected on decryption."],
        ["Shamir's Secret Sharing", "A way to split a secret into n pieces so that any k of them rebuild it, but fewer than k reveal absolutely nothing."],
        ["GF(256)", "A finite field of 256 elements — the mathematical playground where the secret-sharing arithmetic happens (the same one AES uses)."],
        ["Lagrange interpolation", "The math used to rebuild the secret from k shares by fitting a polynomial through them."],
        ["DCT (Discrete Cosine Transform)", "A frequency transform (the heart of JPEG) used here to hide a fingerprint where it survives recompression."],
        ["Watermark fingerprint", "An invisible code identifying which candidate a paper copy belongs to."],
        ["PSNR", "A measure of image quality; ~36 dB here means the watermark is invisible to the eye."],
        ["Hash chain", "A sequence of records each sealed with a SHA-256 hash of the previous one, so any change anywhere is detectable."],
        ["Custodian", "A trusted official who holds one share of the decryption key."],
        ["Time gate", "The server rule that refuses decryption until the official release time, no matter what."],
        ["JWT", "A signed token that proves who a logged-in user is and what role they hold."],
        ["Dynamic assembly", "Building each candidate's paper on demand from a bank of encrypted questions, so no full paper exists before release."],
        ["Blueprint", "The recipe for a dynamic exam: which sections, topics, difficulties, and how many questions each."],
        ["Containment match", "A similarity measure: what fraction of a question's words appear in a leaked text — used to detect text leaks."],
    ],
    widths=[1.8, 4.7])

doc.add_paragraph()
add_hr(doc)
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Trace — built for India's FAR AWAY 2026. Real cryptography, real tracing, real audit. No slideware.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = MUTE

out = "/home/ashok/project/Trace_Project_Documentation.docx"
doc.save(out)
print("Saved:", out)
