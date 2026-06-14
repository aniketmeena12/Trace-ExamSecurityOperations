"""Generate a standalone Word (.docx) demo script for the Trace video.

Run:  backend/.venv/bin/python scripts/build_demo_doc.py
Output: Trace_Demo_Script.docx in the repo root.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK     = RGBColor(0x14, 0x1B, 0x2E)
ACCENT  = RGBColor(0x0E, 0x7A, 0x8C)
ACCENT2 = RGBColor(0x1B, 0xA5, 0xB8)
MUTE    = RGBColor(0x5B, 0x66, 0x77)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)


def shade(target, hex_fill):
    # target may be a table _Cell or a paragraph's lxml <w:p> element.
    if hasattr(target, "_tc"):
        props = target._tc.get_or_add_tcPr()
    else:
        props = target.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    props.append(shd)


def cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def hr(doc, color="1BA5B8"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def beat(doc, time, title, say, show):
    """One timed beat: a shaded time+title bar, the narration, and the screen note."""
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(10)
    bar.paragraph_format.space_after = Pt(2)
    shade(bar._p, "EAF6F8")
    r = bar.add_run(f"  {time}   ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    r2 = bar.add_run(title)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = INK

    say_p = doc.add_paragraph()
    say_p.paragraph_format.left_indent = Inches(0.15)
    lead = say_p.add_run("SAY  ")
    lead.bold = True
    lead.font.size = Pt(9)
    lead.font.color.rgb = ACCENT2
    body = say_p.add_run(say)
    body.font.size = Pt(11)

    show_p = doc.add_paragraph()
    show_p.paragraph_format.left_indent = Inches(0.15)
    lead2 = show_p.add_run("SHOW  ")
    lead2.bold = True
    lead2.font.size = Pt(9)
    lead2.font.color.rgb = MUTE
    body2 = show_p.add_run(show)
    body2.italic = True
    body2.font.size = Pt(10)
    body2.font.color.rgb = MUTE


# =============================================================================
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK

# ---- Title -------------------------------------------------------------------
for _ in range(2):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("TRACE")
r.font.size = Pt(48)
r.bold = True
r.font.color.rgb = ACCENT
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Video Demo Script")
r.font.size = Pt(20)
r.font.color.rgb = INK
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("2–5 minute screen recording with voice-over  ·  target ≈ 3.5 min")
r.font.size = Pt(11)
r.italic = True
r.font.color.rgb = MUTE
doc.add_paragraph()
hr(doc)

intro = doc.add_paragraph()
r = intro.add_run(
    "A judge should understand the whole project from this video alone: the problem, the "
    "solution, the features, and a live demo. Below, "
)
r.font.size = Pt(11)
r = intro.add_run("SAY")
r.bold = True
r.font.color.rgb = ACCENT2
r.font.size = Pt(11)
r = intro.add_run(" is what you narrate; ")
r.font.size = Pt(11)
r = intro.add_run("SHOW")
r.bold = True
r.font.color.rgb = MUTE
r.font.size = Pt(11)
r = intro.add_run(
    " is what's on screen. Before recording: start the backend (:8000) and frontend (:5173), "
    "hard-refresh the browser, and have the role-switcher ready."
)
r.font.size = Pt(11)

# ---- Beats -------------------------------------------------------------------
h = doc.add_paragraph()
r = h.add_run("The Script")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = ACCENT

beat(doc, "0:00–0:20", "Hook",
     "Every year, exam-paper leaks destroy the careers of millions of honest students. One "
     "insider with early access can compromise an entire exam — and once it's out, nobody can "
     "prove whose copy it was. This is Trace: a paper impossible to open before the exam, and "
     "traceable to its exact source if it ever leaks.",
     "Title screen, or the Trace dashboard header (\"Exam Security Ops\").")

beat(doc, "0:20–0:45", "The problem",
     "Two things go wrong in real life. First, the insider leak — someone between the printer "
     "and the exam hall opens the paper early. Second, the accountability gap — even when a leak "
     "is caught, there's no way to trace it back, so no one is ever held responsible.",
     "A simple problem slide, or stay on the dashboard.")

beat(doc, "0:45–1:15", "The solution",
     "Trace solves both with real cryptography. Three guarantees: the paper is encrypted with "
     "AES-256 and the key is split across five custodians with Shamir's Secret Sharing — any "
     "three needed — plus a server-enforced time lock. Every copy carries an invisible watermark "
     "that survives a screen photo. And every action is written to a tamper-evident audit log. "
     "Plus two layers that crush the leak itself: dynamic per-candidate papers and a leak-match "
     "detector.",
     "Slowly scroll the four role chips — Admin, Custodian, Candidate, Investigator.")

beat(doc, "1:15–1:45", "Seal a paper  ·  Admin",
     "As the exam controller I seal a new exam. I choose Dynamic mode — instead of one paper, "
     "the system assembles a unique paper for each candidate from an encrypted question bank. I "
     "set the release to thirty seconds and seal it.",
     "Admin → Seal a New Exam → Dynamic bank toggle → blueprint → +30s → Seal & Distribute. "
     "Point to the vault, the time gate, and the Dynamic Assembly strip.")

beat(doc, "1:45–2:15", "Unlock ceremony  ·  Custodians",
     "Now the custodians. Each holds one key share, masked until the release window. As they "
     "submit, the quorum fills — two of three, three of three. But even with every valid share "
     "in, the vault stays sealed until the clock hits zero. That's the time gate. And it opens.",
     "Custodian → submit a share → switch custodian → submit. Back to Admin ops panel → the "
     "vault animates UNLOCKED.")

beat(doc, "2:15–2:45", "Unique watermarked papers  ·  Candidate",
     "Here's the candidate's own paper, stamped with an invisible watermark unique to them. Now "
     "switch to a different candidate — the questions are different. Every candidate gets a "
     "content-distinct, individually-fingerprinted paper, so a leak exposes only one variant and "
     "points straight at its owner.",
     "Candidate cand001 → watermarked paper + fingerprint. Role-switch to cand002 → visibly "
     "different questions.")

beat(doc, "2:45–3:25", "Catch the leak  ·  Investigator",
     "Now the leak. Suppose the paper turns up on a chat group. I paste the text into the "
     "Leak-Match Detector — it matches the questions against the bank and names the suspect, "
     "R-001. If instead I have a photo, Forensic Trace recovers the watermark even after JPEG "
     "re-compression and identifies the source at 100% confidence. Every detection becomes a "
     "case file with the candidate's full details, and the whole investigation is sealed in a "
     "SHA-256 hash chain — Verify Integrity: intact.",
     "Investigator → Leak-Match (paste text) → suspect. Forensic Trace (upload image) → traced. "
     "Case Files → open the case → full candidate card. Verify Integrity → chain intact.")

beat(doc, "3:25–3:40", "Close",
     "Trace makes early opening cryptographically impossible, every copy traceable, and every "
     "record tamper-proof — backed by 86 passing tests and crypto built from scratch. Leaks "
     "become impossible to hide, and impossible to deny. Thank you.",
     "Return to the dashboard header; optional end card with the three guarantees.")

# ---- Checklist + logins ------------------------------------------------------
doc.add_page_break()
h = doc.add_paragraph()
r = h.add_run("Before You Hit Record")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = ACCENT
for item in [
    "Backend running: cd backend && uvicorn trace.api.app:app  →  http://127.0.0.1:8000",
    "Frontend running: cd frontend && npm run dev  →  http://localhost:5173  (hard-refresh once)",
    "Pre-seed a dynamic exam that is already UNLOCKED so you don't wait on the timer live — "
    "or set release to +20s and edit the pause out.",
    "Have one candidate's paper image downloaded and its text copied, ready to use as the leak.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(11)

h = doc.add_paragraph()
r = h.add_run("Demo Logins")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = ACCENT
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Light Grid Accent 1"
for i, head in enumerate(["Role", "Username", "Password"]):
    cell_text(tbl.rows[0].cells[i], head, bold=True, color=WHITE, size=10)
    shade(tbl.rows[0].cells[i], "0E7A8C")
for role, usr, pwd in [
    ("Admin / Exam Controller", "admin", "admin123"),
    ("Custodians (1–5)", "cust1 … cust5", "custodian1 … custodian5"),
    ("Candidate", "cand001 / cand002", "candidate1 / candidate2"),
    ("Investigator", "investigator", "invest123"),
]:
    cells = tbl.add_row().cells
    cell_text(cells[0], role, size=10)
    cell_text(cells[1], usr, size=10)
    cell_text(cells[2], pwd, size=10)

doc.add_paragraph()
h = doc.add_paragraph()
r = h.add_run("If a judge asks: \"Is the crypto real?\"")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = ACCENT
ans = doc.add_paragraph()
ans.paragraph_format.left_indent = Inches(0.15)
r = ans.add_run(
    "Yes — AES-256-GCM via pycryptodome, but Shamir's Secret Sharing is implemented from scratch "
    "over the GF(256) finite field with its own unit tests; the watermark is a Koch–Zhao DCT "
    "scheme; the audit log is a real SHA-256 hash chain. 86 tests prove it, including that no two "
    "custodians can ever reconstruct the key."
)
r.font.size = Pt(11)
r.italic = True

doc.add_paragraph()
hr(doc)
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Trace — built for India's FAR AWAY 2026. Real cryptography, real tracing, real audit.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = MUTE

out = "/home/ashok/project/Trace_Demo_Script.docx"
doc.save(out)
print("Saved:", out)
